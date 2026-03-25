from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from database import get_db
from models import User, TestResult
from schemas import TestResultCreate, TestResultOut
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, io

router = APIRouter(prefix="/api/tests", tags=["tests"])

# ═══════════════════════════════════════════════════════
# WORD HELPERS
# ═══════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _para(doc, text="", bold=False, size=11, color=None,
          align=None, space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def _heading(doc, text, size=13, space_before=14):
    return _para(doc, text, bold=True, size=size,
                 space_before=space_before, space_after=6)

def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)

def _explain_error(question: str, correct_ans: str) -> str:
    q = question.lower()
    if "сумму" in q or "a + b" in q:
        return "Сложение: складывайте элементы с совпадающими индексами (i,j)."
    if "скаляр" in q or "умножьте на" in q:
        return "Умножение на скаляр: каждый элемент умножается на k."
    if "транспон" in q or "aᵀ" in q:
        return "Транспонирование: строки становятся столбцами. (Aᵀ)ᵢⱼ = Aⱼᵢ."
    if "обратн" in q or "a⁻¹" in q:
        return "A⁻¹ = (1/det)·[[d,−b],[−c,a]]. Сначала det = ad−bc."
    if "2×2" in q:
        return "det(2×2) = a·d − b·c."
    if "3×3" in q:
        return "det(3×3): разложение по строке 1 со знаками +−+."
    if "произведение" in q or "a·b" in q:
        return "c_ij = строка i из A · столбец j из B."
    if "ранг" in q:
        return "Ранг = число ненулевых строк после приведения методом Гаусса."
    if "систему" in q or "x₁" in q:
        return "Крамер: x_i = det(Aᵢ)/det(A)."
    if "размер" in q:
        return "Размер: строки × столбцы."
    if "степень" in q or "^" in q:
        return "A^n = A перемноженная на себя n раз."
    if "минор" in q or "дополнение" in q:
        return "Mᵢⱼ: det без строки i и столбца j. Aᵢⱼ = (−1)^(i+j)·Mᵢⱼ."
    return f"Правильный ответ: {correct_ans}"


def build_word_report(tr, user, detailed) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── ШАПКА ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    r = title_p.add_run("MatriX")
    r.bold = True; r.font.size = Pt(30); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    rs = sub_p.add_run("Результаты тестирования")
    rs.font.size = Pt(11); rs.font.name = "Arial"
    rs.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

    _divider(doc)

    # ── INFO TABLE ──────────────────────────────────────
    _para(doc, space_before=4, space_after=0)
    from docx.shared import Inches as In
    from docx.enum.table import WD_TABLE_ALIGNMENT
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    # left cell — ученик
    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    _set_cell_bg(lc, 'F5F5F5')
    _set_cell_bg(rc, 'FFFFFF')
    lc.width = Inches(3)
    rc.width = Inches(3)

    ts = tr.timestamp.strftime("%d.%m.%Y %H:%M") if tr.timestamp else \
         datetime.now().strftime("%d.%m.%Y %H:%M")

    lc_p = lc.paragraphs[0]
    lc_p.paragraph_format.space_before = Pt(4)
    lc_p.paragraph_format.space_after  = Pt(2)
    def add_cell_line(cell_obj, label, value, first=False):
        p = cell_obj.paragraphs[0] if first else cell_obj.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        rl = p.add_run(label)
        rl.bold = True; rl.font.name = "Arial"; rl.font.size = Pt(10)
        rv = p.add_run(value)
        rv.font.name = "Arial"; rv.font.size = Pt(10)

    if user:
        add_cell_line(lc, "Ученик: ", f"{user.first_name} {user.last_name}", first=True)
        add_cell_line(lc, "Класс: ",  user.grade)
    add_cell_line(lc, "Дата: ", ts, first=(not user))

    add_cell_line(rc, "Тест: ", tr.test_title, first=True)
    time_str = f"{tr.time_spent // 60:02d}:{tr.time_spent % 60:02d}"
    add_cell_line(rc, "Время: ", time_str)
    add_cell_line(rc, "Вопросов: ", str(tr.total_questions))

    _para(doc, space_before=2, space_after=0)
    _divider(doc)

    # ── ИТОГ ────────────────────────────────────────────
    _heading(doc, "Итоговый результат", size=14, space_before=10)

    pct = tr.percentage
    if pct >= 80:
        pct_color = (0x16, 0xA3, 0x4A)
        grade_label = "Отлично"
    elif pct >= 60:
        pct_color = (0x65, 0xA3, 0x0D)
        grade_label = "Хорошо"
    elif pct >= 40:
        pct_color = (0xD9, 0x7B, 0x06)
        grade_label = "Удовлетворительно"
    else:
        pct_color = (0xDC, 0x26, 0x26)
        grade_label = "Неудовлетворительно"

    score_p = doc.add_paragraph()
    score_p.paragraph_format.space_before = Pt(4)
    score_p.paragraph_format.space_after  = Pt(2)
    rp = score_p.add_run(f"{pct}%")
    rp.bold = True; rp.font.size = Pt(36); rp.font.name = "Arial"
    rp.font.color.rgb = RGBColor(*pct_color)
    rs2 = score_p.add_run(f"  —  {grade_label}")
    rs2.bold = True; rs2.font.size = Pt(14); rs2.font.name = "Arial"
    rs2.font.color.rgb = RGBColor(*pct_color)

    _para(doc, f"Правильных ответов: {tr.correct_answers} из {tr.total_questions}",
          size=11, space_before=2, space_after=2)
    _para(doc, f"Затраченное время: {time_str}",
          size=11, space_before=0, space_after=6)

    _divider(doc)

    # ── ПОДРОБНЫЕ ОТВЕТЫ ────────────────────────────────
    _heading(doc, "Разбор ответов", size=14, space_before=10)

    correct_count  = sum(1 for q in detailed if q.get("is_correct"))
    wrong_count    = len(detailed) - correct_count

    stat_p = doc.add_paragraph()
    stat_p.paragraph_format.space_before = Pt(2)
    stat_p.paragraph_format.space_after  = Pt(10)
    rc1 = stat_p.add_run(f"✓ Верно: {correct_count}   ")
    rc1.font.name = "Arial"; rc1.font.size = Pt(11); rc1.bold = True
    rc1.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    rc2 = stat_p.add_run(f"✗ Неверно: {wrong_count}")
    rc2.font.name = "Arial"; rc2.font.size = Pt(11); rc2.bold = True
    rc2.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    for i, q in enumerate(detailed):
        is_correct  = q.get("is_correct", False)
        question    = q.get("question", "").replace("\n", "  |  ")
        user_ans    = q.get("user_answer", "(не отвечено)")
        correct_ans = q.get("correct_answer", "")

        # Номер + статус
        num_p = doc.add_paragraph()
        num_p.paragraph_format.space_before = Pt(8)
        num_p.paragraph_format.space_after  = Pt(3)
        badge = "✓" if is_correct else "✗"
        badge_color = RGBColor(0x16,0xA3,0x4A) if is_correct else RGBColor(0xDC,0x26,0x26)
        rb = num_p.add_run(f"{badge}  ")
        rb.bold = True; rb.font.size = Pt(12); rb.font.name = "Arial"
        rb.font.color.rgb = badge_color
        rn = num_p.add_run(f"Вопрос {i + 1}")
        rn.bold = True; rn.font.size = Pt(12); rn.font.name = "Arial"

        # Текст задания
        _para(doc, question, size=10,
              color=(0x44, 0x44, 0x44), space_before=0, space_after=4)

        # Таблица ответов
        atbl = doc.add_table(rows=2, cols=2)
        atbl.style = 'Table Grid'
        headers = ["Ответ ученика", "Правильный ответ"]
        values  = [user_ans, correct_ans]
        bg_cols = ['FEF2F2', 'F0FDF4'] if not is_correct else ['F0FDF4', 'F0FDF4']

        for col_i in range(2):
            hcell = atbl.cell(0, col_i)
            vcell = atbl.cell(1, col_i)
            _set_cell_bg(hcell, 'F5F5F5')
            _set_cell_bg(vcell, bg_cols[col_i])

            hp = hcell.paragraphs[0]
            hp.paragraph_format.space_before = Pt(2)
            hp.paragraph_format.space_after  = Pt(2)
            hr = hp.add_run(headers[col_i])
            hr.bold = True; hr.font.name = "Arial"; hr.font.size = Pt(9)
            hr.font.color.rgb = RGBColor(0x55,0x55,0x55)

            vp = vcell.paragraphs[0]
            vp.paragraph_format.space_before = Pt(3)
            vp.paragraph_format.space_after  = Pt(3)
            vr = vp.add_run(values[col_i])
            vr.font.name = "Arial"; vr.font.size = Pt(10)
            if col_i == 0 and not is_correct:
                vr.font.color.rgb = RGBColor(0xDC,0x26,0x26)
            elif col_i == 1:
                vr.font.color.rgb = RGBColor(0x16,0xA3,0x4A)
                vr.bold = True

        # Объяснение ошибки — только если неверно
        if not is_correct:
            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.space_before = Pt(4)
            exp_p.paragraph_format.space_after  = Pt(2)
            re1 = exp_p.add_run("Разбор:  ")
            re1.bold = True; re1.font.name = "Arial"; re1.font.size = Pt(10)
            re1.font.color.rgb = RGBColor(0x78,0x35,0x0F)
            re2 = exp_p.add_run(_explain_error(question, correct_ans))
            re2.font.name = "Arial"; re2.font.size = Pt(10); re2.italic = True
            re2.font.color.rgb = RGBColor(0x92,0x40,0x0E)

        _para(doc, space_before=0, space_after=0)

    _divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(8)
    rf = footer_p.add_run(f"Сформировано системой MatriX  ·  {datetime.now().strftime('%d.%m.%Y')}")
    rf.font.name = "Arial"; rf.font.size = Pt(9)
    rf.font.color.rgb = RGBColor(0xA0,0xA0,0xA0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════
# ROUTES
# ═══════════════════════════════════════════════════════

@router.post("/results", response_model=TestResultOut)
async def save_result(data: TestResultCreate, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(User).where(User.id == data.user_id))
    if not res.scalar_one_or_none():
        raise HTTPException(404, "Пользователь не найден")
    tr = TestResult(**data.model_dump())
    db.add(tr)
    await db.commit()
    await db.refresh(tr)
    return tr

@router.get("/results/user/{user_id}", response_model=list[TestResultOut])
async def get_user_results(user_id: int, db: AsyncSession = Depends(get_db)):
    res = await db.execute(
        select(TestResult)
        .where(TestResult.user_id == user_id)
        .order_by(TestResult.timestamp.desc())
    )
    return res.scalars().all()

@router.get("/results/{result_id}/word")
async def download_word(result_id: int, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(TestResult).where(TestResult.id == result_id))
    tr = res.scalar_one_or_none()
    if not tr:
        raise HTTPException(404, "Результат не найден")

    user_res = await db.execute(select(User).where(User.id == tr.user_id))
    user = user_res.scalar_one_or_none()

    try:
        detailed = json.loads(tr.detailed_results)
    except Exception:
        detailed = []

    buf = build_word_report(tr, user, detailed)
    from urllib.parse import quote
    surname  = user.last_name if user else "user"
    filename = f"MatriX_{surname}_{tr.test_title.replace(' ', '_')}.docx"
    encoded  = quote(filename, safe="")

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )